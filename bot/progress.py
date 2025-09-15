import os
import asyncio
import win32print
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import barcode as barcode
from barcode.writer import ImageWriter
import win32com.client
import psutil
import pythoncom

def kill_task():
    task_name = 'WORD.EXE'
    for proc in psutil.process_iter(['pid', 'name']):
        
        if task_name.lower() in proc.info['name'].lower(): 
            try:
                proc.kill()  
            except psutil.NoSuchProcess:
                pass
            except psutil.AccessDenied:
                pass

def generate_barcode(data, filename='barcode.png'):
    """ Shtrixkod yaratish va saqlash """
    barcode_class = barcode.get_barcode_class('ean13')
    barcode_obj = barcode_class(data, writer=ImageWriter())
    return barcode_obj.save(filename, options={
        "module_height": 15.0,
        "font_size": 10,
        "dpi": 650,
        "text_distance": 5.0,
    })

async def update_document(template_path, new_file_path, name, price, usd_price, barcode_image_path):
    """ Hujjatni yangilash va saqlash """
    doc = await asyncio.to_thread(Document, template_path)
    if len(name) > 50:
        name = name[:50]
    for para in doc.paragraphs:
        if "Mahsulotnomi" in para.text:
            para.text = ""
            run = para.add_run(f" {name} / {usd_price}")
            run.font.size = Pt(7)
            run.bold = True
            para.paragraph_format.left_indent = Cm(0.10)
        if "50000" in para.text:
            para.text = ""
            run = para.add_run(f"   {price} сум")
            run.font.size = Pt(14)
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Shtrixkodni qo'shish (agar kerak bo'lsa)
    
    # if os.path.exists(barcode_image_path):
    #     paragraph = doc.add_paragraph()
    #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     run = paragraph.add_run()
    #     run.add_picture(barcode_image_path, width=Cm(4.2), height=Cm(1.5))
    
    os.makedirs("documents", exist_ok=True)
    save_path = os.path.join("documents", new_file_path)
    await asyncio.to_thread(doc.save, save_path)
    return save_path

async def print_document(file_path, copies=1):
    kill_task()
    file_path = fr"C:\Users\alfatech.uz\Documents\narx_chiqarbot\bot\{file_path}"
    if os.path.exists(file_path):
        try:
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, _print_document_sync, file_path, copies)
        except Exception as e:
            print(f"Chop etishda xatolik yuz berdi: {e}")
    else:
        print("Fayl topilmadi!")  

def _print_document_sync(file_path, copies):
    """ Hujjatni sinxron ravishda chop etish """
    try:
        pythoncom.CoInitialize()

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(file_path)

        word.ActivePrinter = "Xprinter XP-303B (копия 1)"

        for _ in range(copies):
            doc.PrintOut()
        doc.Close(False)
        word.Quit()

        print(f"✅ {file_path} fayli {copies} nusxada Xprinter XP-303B (копия 1) printeridan chop etish uchun yuborildi.")
        pythoncom.CoUninitialize()
        return True
    except Exception as e:
        print(f"❌ Xatolik: {e}")
        return False  

async def print_barcode(word_name, data_to_encode, name, price, usd_price, barcode_name='barcode.png', page=1):
    try:
        new_file_path = f'{word_name}.docx'
        barcode_image_path = await asyncio.to_thread(generate_barcode, data_to_encode, barcode_name)
        saved_file_path = await update_document('aaa.docx', new_file_path, name, price, usd_price, barcode_image_path)
        await print_document(saved_file_path, page)
        os.remove(f"{barcode_name}.png")
        os.remove(f"documents/{word_name}.docx")
        return True
    except Exception as e:
        print(e)
        if os.path.exists(f"{barcode_name}.png"):
            os.remove(f"{barcode_name}.png")
        return False